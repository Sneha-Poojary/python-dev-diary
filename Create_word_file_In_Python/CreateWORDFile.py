from docx import Document
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Create the document
doc = Document()

# --- Set default font to Times New Roman ---
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
# Ensure Times New Roman is used for complex scripts too
rPr = style.element.rPr
rFonts = rPr.rFonts
rFonts.set(qn('w:ascii'), 'Times New Roman')
rFonts.set(qn('w:hAnsi'), 'Times New Roman')
rFonts.set(qn('w:eastAsia'), 'Times New Roman')
rFonts.set(qn('w:cs'), 'Times New Roman')

# --- Add a border to the page ---
def add_page_border(section):
    pgBorders = OxmlElement('w:pgBorders')
    pgBorders.set(qn('w:offsetFrom'), 'page')

    for border_name in ['top', 'left', 'bottom', 'right']:
        border_el = OxmlElement(f'w:{border_name}')
        border_el.set(qn('w:val'), 'single')
        border_el.set(qn('w:sz'), '12')     # size
        border_el.set(qn('w:space'), '24')  # spacing from text
        border_el.set(qn('w:color'), '000000')  # black
        pgBorders.append(border_el)

    sectPr = section._sectPr
    sectPr.append(pgBorders)

# Apply border to first (and usually only) section
add_page_border(doc.sections[0])

# Create the DOCX document
#doc = Document()
doc.add_heading('Operator Guide: CAN Signal Publisher Step', 0)

# Overview
doc.add_heading('Overview', level=1)
doc.add_paragraph(
    "This OpenTAP test step is used to validate and publish CAN signal values for VCU calibration. "
    "It interprets a comma-separated CAN payload string, extracts specific bytes based on the selected action type, "
    "and publishes the corresponding signal value to the test result."
)

# Inputs
doc.add_heading('Inputs', level=1)
inputs_table = doc.add_table(rows=1, cols=2)
inputs_table.style = 'Table Grid'
hdr_cells = inputs_table.rows[0].cells
hdr_cells[0].text = 'Parameter'
hdr_cells[1].text = 'Description'
inputs = [
    ("Output1", "Comma-separated CAN frame string (e.g., 0xCFFC827, 01, 02, 03, 04, 05)"),
    ("Actions", "Signal name to extract from the payload (dropdown selection)"),
    ("Logging", "Path to save execution logs"),
]
for param, desc in inputs:
    row = inputs_table.add_row().cells
    row[0].text = param
    row[1].text = desc

# Outputs
doc.add_heading('Outputs', level=1)
outputs_table = doc.add_table(rows=1, cols=2)
outputs_table.style = 'Table Grid'
hdr_cells = outputs_table.rows[0].cells
hdr_cells[0].text = 'Output'
hdr_cells[1].text = 'Description'
outputs = [
    ("RetVal", "Extracted signal value (based on selected action)"),
]
for out, desc in outputs:
    row = outputs_table.add_row().cells
    row[0].text = out
    row[1].text = desc

# How it Works
doc.add_heading('How It Works', level=1)
doc.add_paragraph(
    "1. Reads the input string `Output1` which includes the arbitration ID and signal bytes.\n"
    "2. Matches the arbitration ID to known CAN frames (e.g., 0xCFFC827, 0xCFF1027).\n"
    "3. Based on the selected `Actions` value, extracts the specific byte from the frame.\n"
    "4. Publishes the result and updates the `RetVal` field.\n"
    "5. Logs activity and sets the test verdict accordingly."
)

# Supported Actions
doc.add_heading('Supported Actions & Byte Mapping', level=1)
action_table = doc.add_table(rows=1, cols=3)
action_table.style = 'Table Grid'
hdr_cells = action_table.rows[0].cells
hdr_cells[0].text = 'Arbitration ID'
hdr_cells[1].text = 'Action'
hdr_cells[2].text = 'Byte Position'
actions = [
    ("0xCFFCB27", "OBC_DCDC_VCUCommCtrl", "Byte 1"),
    ("0xCFFC827", "OBC_VCUVoltageSetpt_Ref_V", "Byte 1"),
    ("0xCFFC827", "OBC_VCUCurrSetpt_Ref_A", "Byte 2"),
    ("0xCFFC827", "OBC_VCUOpReqCmd", "Byte 3"),
    ("0xCFFC827", "OBC_VCUCtrlMode", "Byte 4"),
    ("0xCFFC827", "OBC_VCUAcknowledgement", "Byte 5"),
    ("0xCFF1027", "DCDC_OutputVtgSetpt_Ref_V", "Byte 1"),
    ("0xCFF1027", "DCDC_VCUAcknowledgement", "Byte 2"),
    ("0xCFF1027", "DCDC_OutputCurrentSetpt_Ref_A", "Byte 3"),
    ("0xCFF1027", "DCDC_ONOFFReq", "Byte 4"),
]
for arb_id, act, byte in actions:
    row = action_table.add_row().cells
    row[0].text = arb_id
    row[1].text = act
    row[2].text = byte

# Example
doc.add_heading('Example', level=1)
doc.add_paragraph(
    "Input:\n"
    "- Output1: 0xCFFC827, 0A, 1B, 2C, 3D, 4E\n"
    "- Actions: OBC_VCUOpReqCmd\n\n"
    "Processing:\n"
    "- Arbitration ID: 0xCFFC827 matches known frame\n"
    "- Byte 3 (2C) is selected based on action\n"
    "- Result Published: 2C\n"
    "- Verdict: PASS"
)
#
# Troubleshooting
doc.add_heading('Troubleshooting', level=1)
trouble_table = doc.add_table(rows=1, cols=3)
trouble_table.style = 'Table Grid'
hdr_cells = trouble_table.rows[0].cells
hdr_cells[0].text = 'Issue'
hdr_cells[1].text = 'Cause'
hdr_cells[2].text = 'Solution'
issues = [
    ("Invalid Arbitration ID", "ID not matched", "Check if Output1 includes valid ID like 0xCFFC827"),
    ("Incorrect Action", "Action doesn't match expected value", "Select correct action from dropdown"),
    ("List Index Error", "Missing or short payload", "Ensure Output1 includes enough comma-separated values"),
]
for issue, cause, solution in issues:
    row = trouble_table.add_row().cells
    row[0].text = issue
    row[1].text = cause
    row[2].text = solution
#
# Save file
# doc_path = "/mnt/data/Publish_CAN_Signal_Operator_Guide.docx"
# doc.save(doc_path)
# doc_path
doc_path = r"C:\Users\sneha\Downloads\Publish_CAN_Signal_Operator_Guide.docx"
doc.save(doc_path)
